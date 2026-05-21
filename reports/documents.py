from io import BytesIO
from pathlib import Path
from textwrap import wrap
from zipfile import ZIP_DEFLATED, ZipFile

from django.utils import timezone
from docx import Document
from docx.shared import Pt
from PIL import Image, ImageDraw, ImageFont

from orders.models import Order, PaymentTransaction


SELLER = {
    'name': 'ООО «Дорскай»',
    'inn': '7200000000',
    'kpp': '720001001',
    'address': '625000, г. Тюмень, ул. Дизайнерская, 12',
    'bank': 'АО «Демо Банк»',
    'bik': '044525000',
    'account': '40702810000000000001',
    'correspondent': '30101810000000000000',
    'phone': '+7 (3452) 00-00-00',
    'email': 'sales@doorsky.local',
}

DOCUMENT_TYPES = [
    ('order', 'Заказ'),
    ('invoice', 'Счет'),
    ('receipt', 'Чек'),
    ('waybill', 'Накладная'),
    ('act', 'Акт'),
]

DOCUMENT_LABELS = dict(DOCUMENT_TYPES)


def format_money(value):
    return f'{value:,.2f}'.replace(',', ' ')


def format_date(value):
    return timezone.localtime(value).strftime('%d.%m.%Y')


def format_datetime(value):
    return timezone.localtime(value).strftime('%d.%m.%Y %H:%M')


def document_number(order, prefix):
    return f'{prefix}-{order.pk:06d}'


def buyer_title(order):
    if order.customer_type == Order.CUSTOMER_COMPANY and order.company_name:
        parts = [order.company_name]
        if order.company_inn:
            parts.append(f'ИНН {order.company_inn}')
        if order.company_kpp:
            parts.append(f'КПП {order.company_kpp}')
        return ', '.join(parts)
    return order.customer_name


def buyer_address(order):
    return order.company_address or order.delivery_address or 'Адрес не указан'


def document_filename(order, document_type, extension='docx'):
    if document_type not in DOCUMENT_LABELS:
        raise ValueError('Unknown document type.')
    return f'doorsky-{document_type}-{order.pk}.{extension}'


def _new_document():
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    return document


def _add_seller_block(document):
    document.add_paragraph(f'Продавец: {SELLER["name"]}')
    document.add_paragraph(f'ИНН/КПП: {SELLER["inn"]} / {SELLER["kpp"]}')
    document.add_paragraph(f'Адрес: {SELLER["address"]}')
    document.add_paragraph(f'Телефон: {SELLER["phone"]} · Email: {SELLER["email"]}')


def _add_bank_block(document):
    document.add_paragraph(f'Банк: {SELLER["bank"]}')
    document.add_paragraph(f'БИК: {SELLER["bik"]}')
    document.add_paragraph(f'Р/с: {SELLER["account"]}')
    document.add_paragraph(f'К/с: {SELLER["correspondent"]}')


def _add_customer_block(document, order):
    document.add_paragraph(f'Покупатель: {buyer_title(order)}')
    document.add_paragraph(f'Контакт: {order.customer_name}, {order.customer_phone}')
    if order.customer_email:
        document.add_paragraph(f'Email: {order.customer_email}')
    document.add_paragraph(f'Адрес: {buyer_address(order)}')


def _add_items_table(document, order, include_unit=True):
    headers = ['№', 'Артикул', 'Наименование', 'Кол-во']
    if include_unit:
        headers.extend(['Цена', 'Сумма'])
    else:
        headers.append('Сумма')

    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for position, item in enumerate(order.items.select_related('product'), start=1):
        row = table.add_row().cells
        row[0].text = str(position)
        row[1].text = item.product.sku
        row[2].text = item.product.name
        row[3].text = str(item.quantity)
        if include_unit:
            row[4].text = format_money(item.unit_price)
            row[5].text = format_money(item.line_total)
        else:
            row[4].text = format_money(item.line_total)
    return table


def _add_total(document, order):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'Итого: {format_money(order.subtotal)} руб.')
    run.bold = True
    document.add_paragraph('НДС: Без НДС')


def _add_signatures(document, left='Продавец', right='Покупатель'):
    document.add_paragraph()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = f'{left}: __________________ / ____________ /'
    table.rows[0].cells[1].text = f'{right}: __________________ / ____________ /'


def _last_successful_transaction(order):
    return order.payment_transactions.filter(status=PaymentTransaction.STATUS_SUCCEEDED).order_by('-processed_at', '-created_at').first()


def build_order_document(order):
    document = _new_document()
    document.add_heading(f'Заказ DoorSkyStore № {document_number(order, "DSO")}', level=1)
    document.add_paragraph(f'Дата: {format_datetime(order.created_at)}')
    document.add_paragraph(f'Статус заказа: {order.get_status_display()}')
    document.add_paragraph(f'Статус оплаты: {order.get_payment_status_display()}')
    _add_customer_block(document, order)
    _add_items_table(document, order)
    _add_total(document, order)
    if order.comment:
        document.add_paragraph(f'Комментарий: {order.comment}')
    return document


def build_invoice_document(order):
    document = _new_document()
    document.add_heading(f'Счет на оплату № {document_number(order, "INV")} от {format_date(order.created_at)}', level=1)
    _add_seller_block(document)
    _add_bank_block(document)
    _add_customer_block(document, order)
    document.add_paragraph(f'Назначение платежа: оплата заказа DoorSkyStore № {order.pk}')
    document.add_paragraph(f'Способ оплаты: {order.get_payment_method_display()}')
    _add_items_table(document, order)
    _add_total(document, order)
    document.add_paragraph('Счет сформирован автоматически в информационной системе DoorSkyStore.')
    _add_signatures(document)
    return document


def build_receipt_document(order):
    document = _new_document()
    transaction = _last_successful_transaction(order)
    receipt_number = order.payment_reference or document_number(order, 'CHK')
    paid_at = order.paid_at or order.created_at

    document.add_heading(f'Чек оплаты № {receipt_number}', level=1)
    document.add_paragraph('Тип документа: имитационный чек интернет-магазина.')
    document.add_paragraph(f'Дата операции: {format_datetime(paid_at)}')
    document.add_paragraph(f'Продавец: {SELLER["name"]}, ИНН {SELLER["inn"]}')
    document.add_paragraph(f'Покупатель: {buyer_title(order)}')
    document.add_paragraph(f'Способ оплаты: {order.get_payment_method_display()}')
    document.add_paragraph(f'Статус оплаты: {order.get_payment_status_display()}')
    if transaction:
        document.add_paragraph(f'Операция эквайринга: {transaction.reference}')
        document.add_paragraph(f'Провайдер: {transaction.provider}')
    _add_items_table(document, order)
    _add_total(document, order)
    document.add_paragraph('Фискализация не выполнялась: документ создан симулятором оплаты для демонстрационного проекта.')
    return document


def build_waybill_document(order):
    document = _new_document()
    document.add_heading(f'Товарная накладная № {document_number(order, "TGN")} от {format_date(order.created_at)}', level=1)
    document.add_paragraph(f'Грузоотправитель: {SELLER["name"]}, {SELLER["address"]}')
    document.add_paragraph(f'Грузополучатель: {buyer_title(order)}, {buyer_address(order)}')
    document.add_paragraph(f'Основание: заказ № {order.pk} от {format_date(order.created_at)}')
    document.add_paragraph(f'Складской резерв: {"оформлен" if order.stock_reserved else "не активен"}')
    _add_items_table(document, order)
    _add_total(document, order)
    _add_signatures(document, left='Отпустил', right='Получил')
    return document


def build_acceptance_act_document(order):
    document = _new_document()
    document.add_heading(f'Акт приема-передачи № {document_number(order, "ACT")} от {format_date(order.created_at)}', level=1)
    document.add_paragraph(f'{SELLER["name"]} передает, а {buyer_title(order)} принимает товары по заказу № {order.pk}.')
    document.add_paragraph(f'Адрес передачи: {order.delivery_address or buyer_address(order)}')
    _add_items_table(document, order, include_unit=False)
    _add_total(document, order)
    document.add_paragraph('Стороны подтверждают отсутствие претензий по количеству и комплектности на момент передачи.')
    _add_signatures(document, left='Передал', right='Принял')
    return document


def _document_payload(order, document_type):
    transaction = _last_successful_transaction(order)
    payloads = {
        'order': {
            'title': f'Заказ DoorSkyStore № {document_number(order, "DSO")}',
            'meta': [
                f'Дата: {format_datetime(order.created_at)}',
                f'Статус заказа: {order.get_status_display()}',
                f'Статус оплаты: {order.get_payment_status_display()}',
            ],
            'blocks': [
                ('Покупатель', [
                    f'Покупатель: {buyer_title(order)}',
                    f'Контакт: {order.customer_name}, {order.customer_phone}',
                    f'Email: {order.customer_email or "не указан"}',
                    f'Адрес: {buyer_address(order)}',
                ]),
            ],
            'note': order.comment,
            'signatures': [],
            'include_unit': True,
        },
        'invoice': {
            'title': f'Счет на оплату № {document_number(order, "INV")} от {format_date(order.created_at)}',
            'meta': [
                f'Продавец: {SELLER["name"]}',
                f'ИНН/КПП: {SELLER["inn"]} / {SELLER["kpp"]}',
                f'Адрес: {SELLER["address"]}',
                f'Банк: {SELLER["bank"]}, БИК {SELLER["bik"]}',
                f'Р/с: {SELLER["account"]}, К/с: {SELLER["correspondent"]}',
            ],
            'blocks': [
                ('Покупатель', [
                    f'Покупатель: {buyer_title(order)}',
                    f'Контакт: {order.customer_name}, {order.customer_phone}',
                    f'Адрес: {buyer_address(order)}',
                ]),
                ('Оплата', [
                    f'Назначение платежа: оплата заказа DoorSkyStore № {order.pk}',
                    f'Способ оплаты: {order.get_payment_method_display()}',
                ]),
            ],
            'note': 'Счет сформирован автоматически в информационной системе DoorSkyStore.',
            'signatures': ['Продавец', 'Покупатель'],
            'include_unit': True,
        },
        'receipt': {
            'title': f'Чек оплаты № {order.payment_reference or document_number(order, "CHK")}',
            'meta': [
                'Тип документа: имитационный чек интернет-магазина.',
                f'Дата операции: {format_datetime(order.paid_at or order.created_at)}',
                f'Продавец: {SELLER["name"]}, ИНН {SELLER["inn"]}',
                f'Покупатель: {buyer_title(order)}',
                f'Способ оплаты: {order.get_payment_method_display()}',
                f'Статус оплаты: {order.get_payment_status_display()}',
            ],
            'blocks': [
                ('Эквайринг', [
                    f'Операция: {transaction.reference}' if transaction else 'Операция: не проведена',
                    f'Провайдер: {transaction.provider}' if transaction else 'Провайдер: DoorSky Pay Simulator',
                ]),
            ],
            'note': 'Фискализация не выполнялась: документ создан симулятором оплаты для демонстрационного проекта.',
            'signatures': [],
            'include_unit': True,
        },
        'waybill': {
            'title': f'Товарная накладная № {document_number(order, "TGN")} от {format_date(order.created_at)}',
            'meta': [
                f'Грузоотправитель: {SELLER["name"]}, {SELLER["address"]}',
                f'Грузополучатель: {buyer_title(order)}, {buyer_address(order)}',
                f'Основание: заказ № {order.pk} от {format_date(order.created_at)}',
                f'Складской резерв: {"оформлен" if order.stock_reserved else "не активен"}',
            ],
            'blocks': [],
            'note': '',
            'signatures': ['Отпустил', 'Получил'],
            'include_unit': True,
        },
        'act': {
            'title': f'Акт приема-передачи № {document_number(order, "ACT")} от {format_date(order.created_at)}',
            'meta': [
                f'{SELLER["name"]} передает, а {buyer_title(order)} принимает товары по заказу № {order.pk}.',
                f'Адрес передачи: {order.delivery_address or buyer_address(order)}',
            ],
            'blocks': [],
            'note': 'Стороны подтверждают отсутствие претензий по количеству и комплектности на момент передачи.',
            'signatures': ['Передал', 'Принял'],
            'include_unit': False,
        },
    }
    try:
        return payloads[document_type]
    except KeyError as exc:
        raise ValueError('Unknown document type.') from exc


def _font_path():
    candidates = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/System/Library/Fonts/Supplemental/Arial Unicode.ttf',
        '/System/Library/Fonts/Supplemental/Arial.ttf',
        '/Library/Fonts/Arial Unicode.ttf',
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def _font(size, bold=False):
    path = _font_path()
    if not path:
        return ImageFont.load_default(size=size)
    if bold and 'DejaVuSans.ttf' in path:
        bold_path = path.replace('DejaVuSans.ttf', 'DejaVuSans-Bold.ttf')
        if Path(bold_path).exists():
            return ImageFont.truetype(bold_path, size)
    return ImageFont.truetype(path, size)


def _wrap_text(text, font, max_width, draw):
    words = str(text).split()
    if not words:
        return ['']

    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f'{current} {word}'
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    result = []
    for line in lines:
        if draw.textbbox((0, 0), line, font=font)[2] <= max_width:
            result.append(line)
        else:
            result.extend(wrap(line, width=42) or [''])
    return result


def _text_height(lines, font, line_gap=7):
    if not lines:
        return 0
    return len(lines) * (font.size + line_gap)


def _table_rows(order, include_unit=True):
    headers = ['№', 'Артикул', 'Наименование', 'Кол-во']
    if include_unit:
        headers.extend(['Цена', 'Сумма'])
    else:
        headers.append('Сумма')

    rows = [headers]
    for position, item in enumerate(order.items.select_related('product'), start=1):
        row = [str(position), item.product.sku, item.product.name, str(item.quantity)]
        if include_unit:
            row.extend([format_money(item.unit_price), format_money(item.line_total)])
        else:
            row.append(format_money(item.line_total))
        rows.append(row)
    return rows


def build_document_image(order, document_type):
    payload = _document_payload(order, document_type)
    width = 1240
    padding = 70
    title_font = _font(34, bold=True)
    heading_font = _font(21, bold=True)
    body_font = _font(18)
    small_font = _font(16)
    total_font = _font(23, bold=True)

    probe = Image.new('RGB', (width, 200), 'white')
    draw = ImageDraw.Draw(probe)
    max_text_width = width - padding * 2

    height = padding + 52
    height += _text_height(_wrap_text(payload['title'], title_font, max_text_width, draw), title_font, 8) + 24
    height += sum(_text_height(_wrap_text(line, body_font, max_text_width, draw), body_font, 7) for line in payload['meta']) + 24
    for heading, lines in payload['blocks']:
        height += heading_font.size + 20
        height += sum(_text_height(_wrap_text(line, body_font, max_text_width, draw), body_font, 7) for line in lines) + 18

    rows = _table_rows(order, payload['include_unit'])
    name_width = 455 if payload['include_unit'] else 620
    table_widths = [54, 170, name_width, 92, 150]
    if payload['include_unit']:
        table_widths.append(160)
    row_heights = []
    for row in rows:
        name_lines = _wrap_text(row[2], small_font, table_widths[2] - 18, draw)
        row_heights.append(max(42, 18 + _text_height(name_lines, small_font, 5)))
    height += sum(row_heights) + 34
    height += 58
    if payload['note']:
        height += _text_height(_wrap_text(payload['note'], body_font, max_text_width, draw), body_font, 7) + 24
    if payload['signatures']:
        height += 92
    height += padding

    image = Image.new('RGB', (width, max(height, 1754)), '#f7f3ea')
    draw = ImageDraw.Draw(image)

    draw.rectangle((0, 0, width, 22), fill='#1b1b17')
    draw.rectangle((0, 22, width, 32), fill='#d7b36a')
    y = padding

    draw.text((padding, y), 'DoorSkyStore', fill='#1b1b17', font=heading_font)
    draw.text((width - padding - 260, y + 4), 'doorsky.local', fill='#6f695f', font=small_font)
    y += 46

    for line in _wrap_text(payload['title'], title_font, max_text_width, draw):
        draw.text((padding, y), line, fill='#11110f', font=title_font)
        y += title_font.size + 8
    y += 18
    draw.line((padding, y, width - padding, y), fill='#d7b36a', width=3)
    y += 24

    for line in payload['meta']:
        for wrapped in _wrap_text(line, body_font, max_text_width, draw):
            draw.text((padding, y), wrapped, fill='#25231f', font=body_font)
            y += body_font.size + 7
    y += 22

    for heading, lines in payload['blocks']:
        draw.text((padding, y), heading, fill='#11110f', font=heading_font)
        y += heading_font.size + 14
        for line in lines:
            for wrapped in _wrap_text(line, body_font, max_text_width, draw):
                draw.text((padding, y), wrapped, fill='#25231f', font=body_font)
                y += body_font.size + 7
        y += 18

    x = padding
    table_top = y
    table_right = padding + sum(table_widths)
    for row_index, row in enumerate(rows):
        row_height = row_heights[row_index]
        fill = '#1f211d' if row_index == 0 else '#fffdf7'
        draw.rectangle((padding, y, table_right, y + row_height), fill=fill, outline='#c8bc9d')
        x = padding
        for cell_index, cell in enumerate(row):
            draw.line((x, y, x, y + row_height), fill='#c8bc9d', width=1)
            color = '#f7f3ea' if row_index == 0 else '#25231f'
            cell_font = small_font
            wrapped = _wrap_text(cell, cell_font, table_widths[cell_index] - 18, draw)
            cell_y = y + 12
            for wrapped_line in wrapped:
                draw.text((x + 9, cell_y), wrapped_line, fill=color, font=cell_font)
                cell_y += small_font.size + 5
            x += table_widths[cell_index]
        draw.line((table_right, y, table_right, y + row_height), fill='#c8bc9d', width=1)
        y += row_height
    draw.rectangle((padding, table_top, table_right, y), outline='#8d6a38', width=2)
    y += 30

    draw.text((padding, y), f'Итого: {format_money(order.subtotal)} руб.', fill='#11110f', font=total_font)
    y += total_font.size + 12
    draw.text((padding, y), 'НДС: Без НДС', fill='#6f695f', font=body_font)
    y += body_font.size + 24

    if payload['note']:
        for line in _wrap_text(payload['note'], body_font, max_text_width, draw):
            draw.text((padding, y), line, fill='#25231f', font=body_font)
            y += body_font.size + 7
        y += 20

    if payload['signatures']:
        left, right = payload['signatures']
        draw.line((padding, y + 44, padding + 430, y + 44), fill='#25231f', width=2)
        draw.line((width - padding - 430, y + 44, width - padding, y + 44), fill='#25231f', width=2)
        draw.text((padding, y + 54), left, fill='#6f695f', font=small_font)
        draw.text((width - padding - 430, y + 54), right, fill='#6f695f', font=small_font)

    return image


BUILDERS = {
    'order': build_order_document,
    'invoice': build_invoice_document,
    'receipt': build_receipt_document,
    'waybill': build_waybill_document,
    'act': build_acceptance_act_document,
}


def build_document(order, document_type):
    try:
        return BUILDERS[document_type](order)
    except KeyError as exc:
        raise ValueError('Unknown document type.') from exc


def build_document_bytes(order, document_type):
    output = BytesIO()
    build_document(order, document_type).save(output)
    return output.getvalue()


def build_document_png_bytes(order, document_type):
    output = BytesIO()
    build_document_image(order, document_type).save(output, format='PNG', optimize=True)
    return output.getvalue()


def build_document_pdf_bytes(order, document_type):
    output = BytesIO()
    image = build_document_image(order, document_type)
    image.save(output, format='PDF', resolution=150.0)
    return output.getvalue()


def build_documents_zip(order, formats=('pdf',)):
    output = BytesIO()
    with ZipFile(output, mode='w', compression=ZIP_DEFLATED) as archive:
        for document_type, _label in DOCUMENT_TYPES:
            if 'pdf' in formats:
                archive.writestr(
                    document_filename(order, document_type, 'pdf'),
                    build_document_pdf_bytes(order, document_type),
                )
            if 'png' in formats:
                archive.writestr(
                    document_filename(order, document_type, 'png'),
                    build_document_png_bytes(order, document_type),
                )
            if 'docx' in formats:
                archive.writestr(
                    document_filename(order, document_type, 'docx'),
                    build_document_bytes(order, document_type),
                )
    return output.getvalue()
